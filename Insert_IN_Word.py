from docx import Document
from docx.shared import Pt

def convert_bullets_to_text(paragraph):
    """
    Convert paragraphs that have bullet/numbering style to plain text with a bullet character.
    """
    text = paragraph.text.strip()

    bullet_chars = ['•', '-', '*', '–', '—']

    if any(text.startswith(b) for b in bullet_chars):
        new_text = text.lstrip(''.join(bullet_chars)).lstrip()
        return f"• {new_text}"
    else:
        import re
        roman_pattern = re.compile(
            r'^\s*(m{0,4}(cm|cd|d?c{0,3})(xc|xl|l?x{0,3})(ix|iv|v?i{0,3}))\.\s+', re.IGNORECASE)
        new_text = roman_pattern.sub('', text).strip()
        if new_text != text:
            return f"• {new_text}"

    return text


def insert_source_as_plain_text(target_path, source_path, marker_text, output_path):
    target_doc = Document(target_path)
    source_doc = Document(source_path)

    insert_index = None
    for i, para in enumerate(target_doc.paragraphs):
        if marker_text.lower() in para.text.lower():
            insert_index = i
            break

    if insert_index is None:
        print(f"Marker '{marker_text}' not found in target document.")
        return

    insert_pos = insert_index + 1
    for para in source_doc.paragraphs:
        clean_text = convert_bullets_to_text(para)

        new_para = target_doc.add_paragraph(clean_text)

        run = new_para.runs[0] if new_para.runs else new_para.add_run(clean_text)
        font = run.font
        font.name = 'Arial'
        font.size = Pt(12)  # All text font size 12

        # Determine if paragraph is header
        is_header = False
        if len(clean_text) < 40 and not clean_text.startswith('•'):
            is_header = True

        for run_source in para.runs:
            if run_source.bold:
                is_header = True
                break

        if is_header:
            run.bold = True

        target_doc._body._element.remove(new_para._element)
        target_doc._body._element.insert(insert_pos, new_para._element)
        insert_pos += 1

    target_doc.save(output_path)
    print(f"Inserted source content as plain text after '{marker_text}'. Saved as '{output_path}'.")


# === Usage example ===
target_doc_path = r"C:\Bots\RFP_Creation\Input\RFP.docx"
source_doc_path = r"C:\Bots\RFP_Creation\Output\Formatted_AI_Response.docx"
output_doc_path = r"C:\Bots\RFP_Creation\Output\final_output.docx"
marker_text = "Find Me"

insert_source_as_plain_text(target_doc_path, source_doc_path, marker_text, output_doc_path)
