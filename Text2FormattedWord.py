from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_bullet(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

def save_markdown_to_word(md_text, file_path):
    doc = Document()

    lines = md_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith('### '):
            para = doc.add_paragraph(line[4:])
            para.style = 'Heading 3'
        elif line.startswith('## '):
            para = doc.add_paragraph(line[3:])
            para.style = 'Heading 2'
        elif line.startswith('# '):
            para = doc.add_paragraph(line[2:])
            para.style = 'Heading 1'
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(line[2:])
            add_bullet(para)
        else:
            para = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = para.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    doc.save(file_path)
    return f"Saved formatted response to: {file_path}"

def main():
    input_file = r"C:\Bots\RFP_Creation\Input\test.txt"
    output_file = r"C:\Bots\RFP_Creation\Output\Formatted_AI_Response.docx"

    try:
        with open(input_file, "r", encoding="latin-1") as file:
            ai_response = file.read()
    except Exception as e:
        return f"Error reading the input file: {e}"

    return save_markdown_to_word(ai_response, output_file)
